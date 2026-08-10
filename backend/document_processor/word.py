import docx
import os
from loguru import logger

async def extract_text_from_word(file_path: str) -> str:
    logger.info(f"📝 Word belgesi okunuyor: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # python-docx kütüphanesi modern .docx formatını destekler.
    if file_ext == '.doc':
        logger.warning(f"Eski format (.doc) tespit edildi. Lütfen .docx kullanmayı tercih edin: {file_path}")
        # Eski .doc dosyaları için basit bir uyarı dönüyoruz
        return "[Uyarı: Bu sistem modern .docx formatındaki Word dosyalarını desteklemektedir. Lütfen dosyanızı .docx olarak kaydedip tekrar yükleyin.]"

    try:
        # Word dosyasını aç
        doc = docx.Document(file_path)
        full_text = []
        
        # 1. Paragrafları oku
        for para in doc.paragraphs:
            if para.text.strip():  # Sadece boş olmayan satırları al
                full_text.append(para.text.strip())
                
        # 2. Tabloları oku (Word içindeki tablolar veri analizi için önemlidir)
        if doc.tables:
            full_text.append("\n--- Belge İçindeki Tablolar ---")
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
                full_text.append("-------------------------------")
                
        logger.success(f"✅ Word başarıyla okundu: {os.path.basename(file_path)}")
        return "\n".join(full_text)
        
    except Exception as e:
        logger.error(f"Word okuma hatası ({file_path}): {str(e)}")
        return f"[Word Okuma Hatası: Bu dosya işlenirken bir sorun oluştu.]"